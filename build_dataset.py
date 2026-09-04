#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
연금 Agent — RAG 데이터셋 구축 파이프라인
=========================================

PDF / DOCX  →  (텍스트 + 표 보존)  →  섹션 단위 청킹  →  chunks.jsonl

핵심 설계
---------
1. 메타데이터 보존   폴더명의 펀드표준코드(KR5109…), 펀드명, 작성기준일, 페이지,
                    섹션 경로를 청크마다 붙인다. RAG 답변에 출처를 달기 위함.
2. 표 보존           수수료·보수율 표는 이 데이터의 핵심 자산이다.
                    extract_text()로는 좌우가 뭉개지므로 extract_tables()로
                    따로 뽑아 마크다운 표로 저장한다.
3. 반복 헤더 제거     페이지마다 반복되는 펀드명 머리글은 노이즈이므로 자동 탐지 후 제거.
4. 재개(resume)      문서 1개 = 결과 파일 1개. 중단돼도 이미 끝난 문서는 건너뛴다.
5. OCR은 최후수단     텍스트 레이어가 살아있으면 OCR을 부르지 않는다.
                    렌더링은 poppler 없이 pypdfium2로 처리(pdf2image보다 훨씬 빠름).
6. 키 분리           CLOVA OCR 키는 .env에서 읽는다. 코드에 넣지 않는다.

사용법
------
    pip install pdfplumber pypdfium2 python-docx python-dotenv requests tqdm
    cp .env.example .env      # 키 입력 (OCR 안 쓸 거면 비워둬도 됨)

    python build_dataset.py --input . --out ./dataset
    python build_dataset.py --input . --out ./dataset --limit 5      # 샘플 테스트
    python build_dataset.py --input . --out ./dataset --workers 4    # 병렬
    python build_dataset.py --input . --out ./dataset --force        # 전체 재추출
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import sys
import time
import unicodedata
import uuid
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Iterable

# ──────────────────────────────────────────────────────────────────────────
# 설정
# ──────────────────────────────────────────────────────────────────────────

SKIP_DIRS = {"venv", ".venv", "__pycache__", ".git", "node_modules", "dataset", ".ipynb_checkpoints"}

# 팀 내부 문서 — 색인에 들어가면 API 응답의 retrieved_context로 새어 나간다.
# 실측(2026-09-05): 아키텍처-분석.pdf가 18청크로 들어가 평가 문항 3건에서
# 노출됐고 BM25 상위 1위였던 문항이 5개였다. .gitignore는 git 추적만 막고
# 이 스캔은 막지 못하므로 여기서 이름으로 제외한다.
# 비교는 nfc()로 정규화한다 — macOS는 파일명을 NFD로 저장한다.
SKIP_FILES = {"아키텍처-분석.pdf"}
DOC_EXTS = {".pdf", ".docx"}

# 청킹 파라미터 — RAG 검색 품질에 직결된다.
CHUNK_TARGET = 900       # 목표 글자 수
CHUNK_MAX = 1400         # 이보다 크면 무조건 자른다
CHUNK_OVERLAP = 150      # 문맥 유실 방지용 겹침
CHUNK_MIN = 120          # 이보다 짧은 조각은 앞 청크에 붙인다

# 텍스트 레이어가 이보다 빈약하면 스캔 이미지로 보고 OCR 시도
OCR_TRIGGER_CHARS = 20
OCR_DPI = 200
OCR_MAX_EDGE = 2400      # 렌더 이미지 긴 변 상한(px)

TABLE_SETTINGS = {
    "vertical_strategy": "lines",
    "horizontal_strategy": "lines",
    "snap_tolerance": 5,
    "join_tolerance": 5,
    "intersection_tolerance": 5,
}

# 섹션 제목으로 볼 패턴 (투자설명서 / 약관 / 규정류 공통)
HEADING_PATTERNS = [
    re.compile(r"^제\s*\d+\s*[부장절조]\s*[.．]?\s*\S"),        # 제 2 부. / 제3장
    re.compile(r"^[<＜〈【\[]\s*[^>＞〉】\]]{2,40}\s*[>＞〉】\]]\s*$"),  # <요약정보>
    re.compile(r"^\d{1,2}\s*[.．]\s*\S{2,60}$"),                 # 1. 집합투자기구의 명칭
    re.compile(r"^[가-힣]\s*[.．]\s*\S{2,60}$"),                 # 가. 투자목적
    re.compile(r"^\(\s*\d{1,2}\s*\)\s*\S{2,60}$"),               # (2) 종류별 수수료
    re.compile(r"^[①-⑳]\s*\S{2,60}$"),
]

# 메타데이터 추출 정규식
RE_FUND_CODE = re.compile(r"\bKR\d[A-Z0-9]{9,11}\b", re.I)
RE_FUND_NAME = [
    re.compile(r"1\s*[.．]\s*집합투자기구\s*명칭\s*[:：]?\s*(.+?)\s*(?:\n|$)"),
    re.compile(r"이\s*투자설명서는\s*(.+?)에\s*대한"),
    re.compile(r"이\s*요약정보는\s*['\"‘“]?(.+?)['\"’”]?\s*(?:집합투자기구|투자신탁)"),
]
_D = r"(\d{4})\s*(?:년|[.．])\s*(\d{1,2})\s*(?:월|[.．])\s*(\d{1,2})\s*일?"
RE_BASE_DATE = [
    re.compile(r"작성\s*기준일\s*[:：]?\s*" + _D),
    re.compile(r"작성\s*기준일\s*\)?\s*[:：]?\s*" + _D),
    re.compile(r"4\s*[.．]\s*작성\s*기준일\s*[:：]?\s*" + _D),
    re.compile(r"증권신고서\s*효력발생일\s*[:：]?\s*" + _D),   # 마지막 폴백
]
RE_KOFIA_CODE = re.compile(r"금융투자협회\s*펀드코드\s*\n?.*?(\d{5,6})", re.S)


# ──────────────────────────────────────────────────────────────────────────
# 자료구조
# ──────────────────────────────────────────────────────────────────────────

@dataclass
class Block:
    """페이지에서 뽑아낸 원자 단위 — 문단 덩어리 또는 표 1개."""
    page: int
    kind: str            # "text" | "table"
    content: str
    section: str = ""
    ocr: bool = False


@dataclass
class DocResult:
    doc_id: str
    source_path: str
    doc_type: str
    fund_code: str | None = None
    fund_name: str | None = None
    kofia_code: str | None = None
    base_date: str | None = None
    n_pages: int = 0
    n_ocr_pages: int = 0
    blocks: list[Block] = field(default_factory=list)
    error: str | None = None


# ──────────────────────────────────────────────────────────────────────────
# 유틸
# ──────────────────────────────────────────────────────────────────────────

def nfc(s: str) -> str:
    """macOS 파일명은 NFD로 저장된다. 비교/저장 전에 NFC로 통일."""
    return unicodedata.normalize("NFC", s)


def squeeze(s: str) -> str:
    s = s.replace(" ", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def is_heading(line: str) -> bool:
    line = line.strip()
    if not (2 <= len(line) <= 70):
        return False
    if line.endswith(("다.", "니다", "습니다", "됩니다", "합니다")):
        return False
    # '24.12.22', "'21.12.23~'22.12.22" 같은 날짜/기간은 제목이 아니다.
    # (숫자 항목 패턴 `1. 제목`에 잘못 걸린다)
    if re.match(r"^['\"]?[\d.\s~,\-'\"]+$", line):
        return False
    return any(p.match(line) for p in HEADING_PATTERNS)


# ──────────────────────────────────────────────────────────────────────────
# 표 처리
# ──────────────────────────────────────────────────────────────────────────

def clean_grid(tb: list[list[Any]]) -> list[list[str]]:
    """pdfplumber 표 격자에서 빈 행/열과 '레이아웃용 가짜 표'를 걷어낸다.

    투자설명서 PDF는 페이지 전체를 사각형으로 감싸는 경우가 많아서
    extract_tables()가 본문 덩어리를 표로 오인한다. 이런 가짜 표를 그대로
    두면 청크가 파이프 기호로 뒤덮여 검색 품질이 떨어진다.
    → '한 행에 3칸 이상 채워진 행'(dense row)이 2개 미만이면 표로 보지 않는다.
    """
    g = [[(c or "").replace("\n", " ").strip() for c in row] for row in tb]
    if not g:
        return []
    ncol = max(len(r) for r in g)
    g = [r + [""] * (ncol - len(r)) for r in g]

    # 1) 내용이 아예 없는 열 제거
    keep = [i for i in range(ncol) if any(r[i] for r in g)]
    # 2) 값이 한 칸뿐인 열은 셀 병합 잔재 → 제거 (행이 충분할 때만)
    if len(g) >= 5:
        keep = [i for i in keep if sum(1 for r in g if r[i]) > 1]
    if len(keep) < 2:
        return []
    g = [[r[i] for i in keep] for r in g]
    g = [r for r in g if any(r)]

    # 3) 희소한 열 제거: 채워진 비율이 10% 미만이고 열이 많을 때
    if len(g) >= 6 and len(g[0]) > 5:
        keep = [i for i in range(len(g[0]))
                if sum(1 for r in g if r[i]) / len(g) >= 0.10]
        if len(keep) >= 2:
            g = [[r[i] for i in keep] for r in g]
            g = [r for r in g if any(r)]

    # 4) 진짜 표인지 판정
    dense = sum(1 for r in g if sum(1 for c in r if c) >= 3)
    if dense < 2:
        return []
    return g


def grid_to_markdown(g: list[list[str]]) -> str | None:
    """표를 마크다운으로. 헤더가 여러 줄인 경우가 많아 1행만 헤더로 쓴다."""
    if len(g) < 2 or len(g[0]) < 2:
        return None
    w = len(g[0])
    rows = ["| " + " | ".join(c or " " for c in g[0]) + " |", "|" + "---|" * w]
    for r in g[1:]:
        rows.append("| " + " | ".join(c or " " for c in r) + " |")
    return "\n".join(rows)


def split_big_table(md: str, max_chars: int = 2500) -> list[str]:
    """긴 표는 헤더를 반복하며 조각낸다 (표 한가운데를 자르지 않기 위해)."""
    if len(md) <= max_chars:
        return [md]
    lines = md.split("\n")
    header, sep, body = lines[0], lines[1], lines[2:]
    out, cur = [], []
    size = len(header) + len(sep)
    for ln in body:
        if cur and size + len(ln) > max_chars:
            out.append("\n".join([header, sep] + cur))
            cur, size = [], len(header) + len(sep)
        cur.append(ln)
        size += len(ln) + 1
    if cur:
        out.append("\n".join([header, sep] + cur))
    return out


# ──────────────────────────────────────────────────────────────────────────
# OCR (CLOVA)
# ──────────────────────────────────────────────────────────────────────────

class ClovaOCR:
    def __init__(self) -> None:
        self.url = os.environ.get("CLOVA_OCR_URL", "").strip()
        self.key = os.environ.get("CLOVA_OCR_SECRET", "").strip()
        self.enabled = bool(self.url and self.key)
        self._session = None

    @staticmethod
    def diagnose() -> str:
        """왜 OCR이 꺼졌는지 사람이 읽을 수 있게 설명한다."""
        url = os.environ.get("CLOVA_OCR_URL", "").strip()
        key = os.environ.get("CLOVA_OCR_SECRET", "").strip()
        if url and key:
            if not url.startswith("http"):
                return f"CLOVA_OCR_URL이 http로 시작하지 않습니다: {url[:40]!r}"
            return ""
        if not Path(".env").exists():
            return (".env 파일이 없습니다. "
                    "`cp .env.example .env` 후 URL과 Secret Key를 채우세요.")
        missing = [n for n, v in (("CLOVA_OCR_URL", url), ("CLOVA_OCR_SECRET", key)) if not v]
        return f".env는 있지만 값이 비어 있습니다: {', '.join(missing)}"

    def __call__(self, image_bytes: bytes) -> str:
        if not self.enabled:
            return ""
        import requests
        if self._session is None:
            self._session = requests.Session()
        payload = {"message": json.dumps({
            "images": [{"format": "jpg", "name": "page"}],
            "requestId": str(uuid.uuid4()),
            "version": "V2",
            "timestamp": int(time.time() * 1000),
        })}
        files = [("file", ("page.jpg", image_bytes, "image/jpeg"))]
        for attempt in range(3):
            try:
                r = self._session.post(self.url, headers={"X-OCR-SECRET": self.key},
                                       data=payload, files=files, timeout=60)
                if r.status_code == 429 or r.status_code >= 500:
                    time.sleep(2 ** attempt)
                    continue
                js = r.json()
                words = []
                for img in js.get("images", []):
                    for f in img.get("fields", []):
                        words.append(f.get("inferText", ""))
                        if f.get("lineBreak"):
                            words.append("\n")
                return squeeze(" ".join(words))
            except Exception as e:  # noqa: BLE001
                if attempt == 2:
                    print(f"    [OCR 실패] {e}", file=sys.stderr)
                time.sleep(2 ** attempt)
        return ""


# ──────────────────────────────────────────────────────────────────────────
# 추출
# ──────────────────────────────────────────────────────────────────────────

def detect_running_headers(page_texts: list[str], min_ratio: float = 0.5) -> set[str]:
    """페이지 절반 이상에서 첫/끝 줄로 반복되는 문자열 = 머리글/꼬리말."""
    from collections import Counter
    c: Counter[str] = Counter()
    for t in page_texts:
        lines = [l.strip() for l in t.split("\n") if l.strip()]
        for l in lines[:2] + lines[-2:]:
            if 4 <= len(l) <= 60:
                c[l] += 1
    n = max(1, len(page_texts))
    return {k for k, v in c.items() if v / n >= min_ratio and n >= 4}


def extract_pdf(path: Path, ocr: ClovaOCR) -> tuple[list[Block], int, int, str]:
    import pdfplumber

    blocks: list[Block] = []
    n_ocr = 0
    raw_pages: list[str] = []
    page_tables: list[list[str]] = []

    with pdfplumber.open(str(path)) as pdf:
        n_pages = len(pdf.pages)
        for pg in pdf.pages:
            raw_pages.append(pg.extract_text() or "")
            mds = []
            try:
                for tb in pg.extract_tables(TABLE_SETTINGS) or []:
                    g = clean_grid(tb)
                    md = grid_to_markdown(g)
                    if md:
                        mds.append(md)
            except Exception:  # noqa: BLE001
                pass
            page_tables.append(mds)

    headers = detect_running_headers(raw_pages)

    pdfium_doc = None
    for i, raw in enumerate(raw_pages):
        page_no = i + 1
        text = raw
        used_ocr = False

        if len(text.strip()) < OCR_TRIGGER_CHARS and ocr.enabled:
            if pdfium_doc is None:
                import pypdfium2 as pdfium
                pdfium_doc = pdfium.PdfDocument(str(path))
            try:
                pil = pdfium_doc[i].render(scale=OCR_DPI / 72).to_pil().convert("RGB")
                # 긴 변을 제한해 전송량·응답시간을 줄인다 (인식률 손해는 거의 없음)
                if max(pil.size) > OCR_MAX_EDGE:
                    r = OCR_MAX_EDGE / max(pil.size)
                    pil = pil.resize((int(pil.width * r), int(pil.height * r)))
                buf = io.BytesIO()
                pil.save(buf, format="JPEG", quality=88)
                text = ocr(buf.getvalue())
                used_ocr = bool(text)
                n_ocr += int(used_ocr)
            except Exception as e:  # noqa: BLE001
                print(f"    [렌더 실패] {path.name} p{page_no}: {e}", file=sys.stderr)

        lines = [l for l in text.split("\n") if l.strip() and l.strip() not in headers]
        body = squeeze("\n".join(lines))
        if body:
            blocks.append(Block(page=page_no, kind="text", content=body, ocr=used_ocr))
        for md in page_tables[i]:
            for piece in split_big_table(md):
                blocks.append(Block(page=page_no, kind="table", content=piece))

    if pdfium_doc is not None:
        pdfium_doc.close()

    return blocks, n_pages, n_ocr, "\n".join(raw_pages[:4])


def extract_docx(path: Path) -> tuple[list[Block], int, int, str]:
    import docx

    d = docx.Document(str(path))
    blocks: list[Block] = []
    buf: list[str] = []
    section = ""

    def flush() -> None:
        nonlocal buf
        if buf:
            blocks.append(Block(page=1, kind="text", content=squeeze("\n".join(buf)), section=section))
            buf = []

    for para in d.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading") or is_heading(t):
            flush()
            section = t
            buf.append(t)
        else:
            buf.append(t)
    flush()

    for tb in d.tables:
        g = clean_grid([[c.text for c in row.cells] for row in tb.rows])
        md = grid_to_markdown(g)
        if md:
            for piece in split_big_table(md):
                blocks.append(Block(page=1, kind="table", content=piece, section=section))

    head = "\n".join(p.text for p in d.paragraphs[:60])
    return blocks, 1, 0, head


# ──────────────────────────────────────────────────────────────────────────
# 메타데이터
# ──────────────────────────────────────────────────────────────────────────

def infer_metadata(path: Path, root: Path, head_text: str) -> dict[str, Any]:
    rel = nfc(str(path.relative_to(root)))
    parts = rel.split(os.sep)

    doc_type = "기타"
    low = rel
    if "투자설명서" in low:
        doc_type = "투자설명서"
    elif "약관" in low:
        doc_type = "약관"
    elif "docs_renamed" in low:
        doc_type = "연금문서"

    # 펀드 표준코드: 폴더명 > 파일명 > 본문 순
    code = None
    for cand in reversed(parts):
        m = RE_FUND_CODE.search(nfc(cand))
        if m:
            code = m.group(0).upper()
            break
    if not code:
        m = RE_FUND_CODE.search(head_text)
        code = m.group(0).upper() if m else None

    name = None
    for pat in RE_FUND_NAME:
        m = pat.search(head_text)
        if m:
            name = squeeze(m.group(1))[:120]
            break

    base_date = None
    for pat in RE_BASE_DATE:
        m = pat.search(head_text)
        if m:
            base_date = f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
            break

    kofia = None
    m = RE_KOFIA_CODE.search(head_text)
    if m:
        kofia = m.group(1)

    doc_id = code or re.sub(r"[^\w가-힣]+", "_", nfc(path.stem))[:80]
    return dict(doc_id=doc_id, doc_type=doc_type, fund_code=code,
                fund_name=name, kofia_code=kofia, base_date=base_date,
                source_path=rel)


# ──────────────────────────────────────────────────────────────────────────
# 섹션 부여 + 청킹
# ──────────────────────────────────────────────────────────────────────────

def assign_sections(blocks: list[Block]) -> None:
    """텍스트 블록의 줄을 훑어 가장 최근 제목을 섹션으로 붙인다."""
    current = ""
    for b in blocks:
        if b.kind == "text":
            found = current
            for line in b.content.split("\n"):
                if is_heading(line):
                    found = line.strip()
            b.section = current or found
            current = found
        else:
            b.section = current


def split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[.。!?])\s+|(?<=다\.)\s*\n|\n", text)
    return [p for p in (x.strip() for x in parts) if p]


def chunk_blocks(blocks: list[Block]) -> list[dict[str, Any]]:
    """표는 통째로 1청크. 텍스트는 섹션 단위로 모아 목표 길이로 자른다."""
    chunks: list[dict[str, Any]] = []

    def emit(page: int, section: str, kind: str, text: str, ocr: bool) -> None:
        text = text.strip()
        if len(text) < 30:
            return
        chunks.append(dict(page=page, section=section, kind=kind, text=text, ocr=ocr))

    buf: list[str] = []
    buf_len = 0
    buf_page = 1
    buf_sec = ""
    buf_ocr = False

    def flush() -> None:
        nonlocal buf, buf_len
        if buf:
            emit(buf_page, buf_sec, "text", "\n".join(buf), buf_ocr)
            buf, buf_len = [], 0

    for b in blocks:
        if b.kind == "table":
            flush()
            emit(b.page, b.section, "table", b.content, False)
            continue

        if b.section != buf_sec:
            flush()
            buf_sec, buf_page, buf_ocr = b.section, b.page, b.ocr
        if not buf:
            buf_page, buf_ocr = b.page, b.ocr

        for sent in split_sentences(b.content):
            if buf_len + len(sent) > CHUNK_MAX and buf:
                emit(buf_page, buf_sec, "text", "\n".join(buf), buf_ocr)
                tail, tlen = [], 0
                for s in reversed(buf):          # 겹침 확보
                    if tlen + len(s) > CHUNK_OVERLAP:
                        break
                    tail.insert(0, s)
                    tlen += len(s)
                buf, buf_len = tail, tlen
                buf_page = b.page
            buf.append(sent)
            buf_len += len(sent) + 1
            if buf_len >= CHUNK_TARGET:
                emit(buf_page, buf_sec, "text", "\n".join(buf), buf_ocr)
                tail, tlen = [], 0
                for s in reversed(buf):
                    if tlen + len(s) > CHUNK_OVERLAP:
                        break
                    tail.insert(0, s)
                    tlen += len(s)
                buf, buf_len = tail, tlen
                buf_page = b.page
    flush()
    return chunks


# ──────────────────────────────────────────────────────────────────────────
# 문서 1개 처리 (워커 진입점)
# ──────────────────────────────────────────────────────────────────────────

def process_one(args: tuple[str, str, str, bool]) -> dict[str, Any]:
    path_s, root_s, outdir_s, force = args
    path, root, outdir = Path(path_s), Path(root_s), Path(outdir_s)
    docs_dir = outdir / "docs"
    docs_dir.mkdir(parents=True, exist_ok=True)

    stem = re.sub(r"[^\w가-힣.-]+", "_", nfc(str(path.relative_to(root))))[:150]
    cache = docs_dir / f"{stem}.json"
    if cache.exists() and not force and cache.stat().st_mtime >= path.stat().st_mtime:
        return {"status": "skip", "cache": str(cache)}

    ocr = ClovaOCR()
    try:
        if path.suffix.lower() == ".pdf":
            blocks, n_pages, n_ocr, head = extract_pdf(path, ocr)
        else:
            blocks, n_pages, n_ocr, head = extract_docx(path)
    except Exception as e:  # noqa: BLE001
        res = DocResult(doc_id=nfc(path.stem), source_path=nfc(str(path.relative_to(root))),
                        doc_type="?", error=f"{type(e).__name__}: {e}")
        cache.write_text(json.dumps(asdict(res), ensure_ascii=False), encoding="utf-8")
        return {"status": "error", "file": nfc(path.name), "error": str(e)}

    meta = infer_metadata(path, root, head)
    assign_sections(blocks)

    res = DocResult(n_pages=n_pages, n_ocr_pages=n_ocr, blocks=blocks, **meta)
    payload = asdict(res)
    payload["chunks"] = chunk_blocks(blocks)
    payload.pop("blocks")
    cache.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    return {"status": "ok", "file": nfc(path.name), "doc_id": meta["doc_id"],
            "pages": n_pages, "ocr": n_ocr, "chunks": len(payload["chunks"])}


# ──────────────────────────────────────────────────────────────────────────
# 메인
# ──────────────────────────────────────────────────────────────────────────

def find_docs(root: Path) -> list[Path]:
    out = []
    for p in root.rglob("*"):
        if any(part in SKIP_DIRS for part in p.parts):
            continue
        if p.is_file() and p.suffix.lower() in DOC_EXTS and not p.name.startswith((".", "~$")):
            if nfc(p.name) in {nfc(x) for x in SKIP_FILES}:
                continue
            out.append(p)
    return sorted(out, key=lambda x: nfc(str(x)))


def merge(outdir: Path) -> tuple[int, int]:
    """문서별 결과를 chunks.jsonl + manifest.csv로 합친다."""
    import csv

    docs_dir = outdir / "docs"
    jsonl = outdir / "chunks.jsonl"
    manifest = outdir / "manifest.csv"

    n_chunks = n_docs = 0
    with jsonl.open("w", encoding="utf-8") as fj, manifest.open("w", encoding="utf-8-sig", newline="") as fm:
        w = csv.writer(fm)
        w.writerow(["doc_id", "doc_type", "fund_code", "fund_name", "kofia_code",
                    "base_date", "n_pages", "n_ocr_pages", "n_chunks", "n_tables",
                    "status", "source_path", "error"])
        for f in sorted(docs_dir.glob("*.json")):
            d = json.loads(f.read_text(encoding="utf-8"))
            chunks = d.get("chunks", [])
            n_docs += 1
            n_tables = sum(1 for c in chunks if c["kind"] == "table")
            pages = d.get("n_pages") or 0
            if d.get("error"):
                status = "ERROR"
            elif not chunks:
                # 텍스트 레이어가 없는 스캔 PDF일 가능성이 높다 → OCR 필요
                status = "NEEDS_OCR" if pages else "EMPTY"
            elif pages and len(chunks) / pages < 0.4:
                status = "LOW_YIELD"          # 페이지 수 대비 추출량이 적음 → 육안 확인
            else:
                status = "OK"
            w.writerow([d.get("doc_id"), d.get("doc_type"), d.get("fund_code"),
                        d.get("fund_name"), d.get("kofia_code"), d.get("base_date"),
                        pages, d.get("n_ocr_pages"), len(chunks), n_tables,
                        status, d.get("source_path"), d.get("error") or ""])
            for i, c in enumerate(chunks):
                rec = {
                    "chunk_id": f"{d['doc_id']}_p{c['page']}_{i:04d}",
                    "doc_id": d.get("doc_id"),
                    "doc_type": d.get("doc_type"),
                    "fund_code": d.get("fund_code"),
                    "fund_name": d.get("fund_name"),
                    "kofia_code": d.get("kofia_code"),
                    "base_date": d.get("base_date"),
                    "page": c["page"],
                    "section": c["section"],
                    "content_type": c["kind"],
                    "is_ocr": c["ocr"],
                    "text": c["text"],
                    "n_chars": len(c["text"]),
                    "source_path": d.get("source_path"),
                }
                fj.write(json.dumps(rec, ensure_ascii=False) + "\n")
                n_chunks += 1
    return n_docs, n_chunks


def main() -> None:
    ap = argparse.ArgumentParser(description="연금 Agent RAG 데이터셋 빌더")
    ap.add_argument("--input", default=".", help="문서 루트 폴더")
    ap.add_argument("--out", default="./dataset", help="결과 폴더")
    ap.add_argument("--workers", type=int, default=1, help="병렬 프로세스 수")
    ap.add_argument("--limit", type=int, default=0, help="처음 N개만 (테스트용)")
    ap.add_argument("--force", action="store_true", help="캐시 무시하고 전체 재추출")
    ap.add_argument("--merge-only", action="store_true", help="추출 없이 병합만")
    a = ap.parse_args()

    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass

    root = Path(a.input).resolve()
    outdir = Path(a.out).resolve()
    outdir.mkdir(parents=True, exist_ok=True)

    if a.merge_only:
        nd, nc = merge(outdir)
        print(f"병합 완료: 문서 {nd}개 / 청크 {nc}개 → {outdir/'chunks.jsonl'}")
        return

    files = find_docs(root)
    if a.limit:
        files = files[:a.limit]
    if not files:
        print(f"'{root}' 아래에 PDF/DOCX가 없습니다.")
        return

    why = ClovaOCR.diagnose()
    ocr_on = not why
    print(f"대상 문서 {len(files)}개 | OCR 폴백: {'ON' if ocr_on else 'OFF'} | 워커 {a.workers}")
    if why:
        print(f"  ⚠️  OCR 비활성: {why}")
        print("      → 스캔 PDF는 추출 결과가 비게 됩니다 (manifest.csv에 NEEDS_OCR로 표시).")

    tasks = [(str(f), str(root), str(outdir), a.force) for f in files]
    done = ok = skipped = errs = 0
    t0 = time.time()

    def report(r: dict[str, Any]) -> None:
        nonlocal ok, skipped, errs
        if r["status"] == "ok":
            ok += 1
            print(f"  [{done}/{len(tasks)}] ✓ {r['file']}  ({r['doc_id']}, {r['pages']}p, "
                  f"청크 {r['chunks']}{', OCR ' + str(r['ocr']) + 'p' if r['ocr'] else ''})")
        elif r["status"] == "skip":
            skipped += 1
        else:
            errs += 1
            print(f"  [{done}/{len(tasks)}] ✗ {r.get('file')}: {r.get('error')}", file=sys.stderr)

    if a.workers > 1:
        with ProcessPoolExecutor(max_workers=a.workers) as ex:
            futs = {ex.submit(process_one, t): t for t in tasks}
            for fu in as_completed(futs):
                done += 1
                report(fu.result())
    else:
        for t in tasks:
            done += 1
            report(process_one(t))

    nd, nc = merge(outdir)
    print(f"\n완료: 성공 {ok} / 건너뜀 {skipped} / 실패 {errs}  ({time.time()-t0:.1f}초)")
    print(f"  {outdir/'chunks.jsonl'}   (문서 {nd} · 청크 {nc})")
    print(f"  {outdir/'manifest.csv'}")

    # 상태별 요약 — 눈으로 바로 다음 할 일을 알 수 있게
    import csv as _csv
    from collections import Counter
    with (outdir / "manifest.csv").open(encoding="utf-8-sig") as f:
        rows = list(_csv.DictReader(f))
    tally = Counter(r["status"] for r in rows)
    print("\n[문서 상태]")
    for st in ("OK", "LOW_YIELD", "NEEDS_OCR", "EMPTY", "ERROR"):
        if tally.get(st):
            print(f"  {st:<10} {tally[st]:>4}건")
    need = [r["source_path"] for r in rows if r["status"] in ("NEEDS_OCR", "EMPTY")]
    if need:
        print(f"\n스캔 PDF로 보이는 문서 {len(need)}건 — OCR 없이는 추출 불가:")
        for p in need[:10]:
            print(f"  · {p}")
        if len(need) > 10:
            print(f"  · … 외 {len(need)-10}건")
        if why:
            print(f"  → {why}")
    blank = [r["doc_id"] for r in rows if r["doc_type"] == "투자설명서" and not r["fund_name"]]
    if blank:
        print(f"\n펀드명 추출 실패 {len(blank)}건: {', '.join(blank[:8])}"
              f"{' …' if len(blank) > 8 else ''}")


if __name__ == "__main__":
    main()
